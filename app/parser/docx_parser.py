"""
DOCX Import Parser
------------------
Reads .docx files using python-docx and converts the document structure into
IAS Block objects, preserving:
  - Heading levels (from Word heading styles Heading 1–6)
  - Bold/italic runs → inline markdown
  - Numbered and bulleted lists
  - Tables
  - Highlighted / colored runs → callouts
  - ALL-CAPS short paragraphs → headings
  - Inline callout patterns from paragraph text
"""
import re
import uuid
from typing import List, Optional

from app.models.document import Block, BlockType, Section, Document


def _is_grayscale(hexcolor: str) -> bool:
    """
    True for near-black/white/gray colors -- the default text color in
    virtually every Word document, which is emphatically not the same
    signal as a run someone deliberately colored on purpose. Same
    threshold as theme_extractor._is_grayscale, kept as a local copy here
    rather than a cross-package import since parser/ and themes/ are
    otherwise independent of each other.
    """
    try:
        h = hexcolor.lstrip("#")
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return max(r, g, b) - min(r, g, b) < 18
    except Exception:
        return True


def _strip_color_marker(text: str) -> str:
    """Remove {{c:HEXCODE}}...{{/c}} markers, keeping the wrapped text --
    for regex classification, which needs to see "Warning: ..." right
    after the ** rather than "{{c:8C2F4B}}Warning: {{/c}} ..."."""
    text = re.sub(r"\{\{c:[0-9A-Fa-f]{6}\}\}", "", text)
    return text.replace("{{/c}}", "")


# Word built-in heading style names
HEADING_STYLES = {
    "heading 1": 1, "heading 2": 2, "heading 3": 3,
    "heading 4": 4, "heading 5": 5, "heading 6": 6,
    "title": 1, "subtitle": 2,
}

# Word list style names (approximate)
LIST_STYLES = {
    "list bullet", "list bullet 2", "list bullet 3",
    "list number", "list number 2", "list number 3",
    "list paragraph",
}

CALLOUT_RE = re.compile(
    r"^(?:\*\*)?(?:Important|Critical|Note|Key Rule)(?:\*\*)?\s*[:\-]\s*(.+)",
    re.IGNORECASE
)
WARNING_RE = re.compile(
    r"^(?:\*\*)?(?:Warning|Danger|Alert|Caution)(?:\*\*)?\s*[:\-]\s*(.+)",
    re.IGNORECASE
)
RELATIONSHIP_RE = re.compile(
    r"^(?:Relationship|Allied with|Enemy of|Mentor to|Rival of|Friend of|Partner of)\s*[:\-]?\s*(.+)",
    re.IGNORECASE
)
TIMELINE_RE = re.compile(
    r"^(?:Year\s+\d{3,4}|\d{4})\s*[:\-]\s*(.+)",
    re.IGNORECASE
)
DEFINITION_RE = re.compile(
    r"^(?:Definition|Term|Glossary)\s*[:\-]\s*(.+)",
    re.IGNORECASE
)


class DocxParser:
    """
    Full structural DOCX parser.
    Falls back gracefully if python-docx is not available.
    """

    def parse(self, filepath: str, source_name: str = "") -> Document:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise RuntimeError("python-docx is required: pip install python-docx")

        raw_doc = DocxDocument(filepath)
        source_name = source_name or filepath

        blocks: List[Block] = []
        counter = [0]
        image_counter = [0]

        def make_id() -> str:
            counter[0] += 1
            return f"b{counter[0]:04d}"

        # ── Walk all block-level elements in document order ──────────────
        # We process paragraphs and tables from the body XML to preserve order.
        body = raw_doc.element.body

        pending_list_items: List[str] = []
        pending_list_style: str = "bullet"

        def flush_list():
            if pending_list_items:
                blocks.append(Block(
                    id=make_id(),
                    type=BlockType.LIST,
                    content="\n".join(pending_list_items),
                    metadata={"list_style": pending_list_style}
                ))
                pending_list_items.clear()

        for child in body.iterchildren():
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            # ── Paragraph ───────────────────────────────────────────────
            if tag == "p":
                from docx.text.paragraph import Paragraph
                para = Paragraph(child, raw_doc)
                style_name = (para.style.name or "").lower() if para.style else ""
                text = self._extract_runs(para)

                image_blocks = self._extract_images(para, raw_doc, make_id, image_counter)
                if image_blocks:
                    flush_list()
                    blocks.extend(image_blocks)

                if not text.strip():
                    flush_list()
                    continue

                # Heading styles
                if style_name in HEADING_STYLES:
                    flush_list()
                    level = HEADING_STYLES[style_name]
                    clean = self._strip_inline_md(text)
                    btype = self._maybe_entity(clean, level)
                    blocks.append(Block(
                        id=make_id(), type=btype,
                        content=clean, level=level,
                        metadata={"style": style_name}
                    ))
                    continue

                # Outline level from paragraph XML (Word outline numbering)
                outline_level = self._get_outline_level(para)
                if outline_level is not None:
                    flush_list()
                    level = outline_level + 1  # 0-based → 1-based
                    clean = self._strip_inline_md(text)
                    btype = self._maybe_entity(clean, level)
                    blocks.append(Block(
                        id=make_id(), type=btype,
                        content=clean, level=level,
                        metadata={"outline_level": outline_level}
                    ))
                    continue

                # List styles
                if style_name in LIST_STYLES or self._is_list_paragraph(para):
                    is_numbered = "number" in style_name
                    pending_list_style = "numbered" if is_numbered else "bullet"
                    pending_list_items.append(text.strip())
                    continue

                # Regular paragraph — flush any pending list first
                flush_list()

                # Check for semantic patterns in text
                block = self._classify_paragraph(text.strip(), make_id)
                if block:
                    blocks.append(block)
                    continue

                # ALL-CAPS short line → heading
                plain = re.sub(r"\*+", "", text).strip()
                if plain.isupper() and len(plain) < 60 and len(plain.split()) <= 6:
                    flush_list()
                    blocks.append(Block(
                        id=make_id(), type=BlockType.HEADING,
                        content=plain.title(), level=2
                    ))
                    continue

                # Check if highlighted (Word highlight → callout)
                if self._has_highlight(para):
                    blocks.append(Block(
                        id=make_id(), type=BlockType.CALLOUT,
                        content=text.strip(),
                        metadata={"source": "highlight"}
                    ))
                    continue

                # Default paragraph
                blocks.append(Block(
                    id=make_id(), type=BlockType.PARAGRAPH,
                    content=text.strip()
                ))

            # ── Table ────────────────────────────────────────────────────
            elif tag == "tbl":
                flush_list()
                from docx.table import Table
                tbl = Table(child, raw_doc)
                table_text = self._parse_table(tbl)
                if table_text:
                    blocks.append(Block(
                        id=make_id(), type=BlockType.TABLE,
                        content=table_text
                    ))

        flush_list()

        title = self._extract_title(blocks, filepath)
        sections = self._build_sections(blocks)

        return Document(
            id=str(uuid.uuid4())[:8],
            title=title,
            all_blocks=blocks,
            sections=sections,
            metadata={
                "source": source_name,
                "paragraph_count": len(raw_doc.paragraphs),
                "parser": "docx",
            }
        )

    # ── Helpers ──────────────────────────────────────────────────────────

    def _extract_images(self, para, raw_doc, make_id, image_counter) -> List[Block]:
        """Find every embedded image in this paragraph's runs, pull its
        bytes via the document's relationship parts, run it through the
        (optional) captioning/OCR step, and return one Block per image.
        Never raises -- a malformed or unreadable image is just skipped
        rather than failing the whole document parse."""
        from app.parser.image_pipeline import process_image, fallback_caption

        ns = {
            "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        }
        out: List[Block] = []
        try:
            blips = para._p.findall(".//a:blip", ns)
        except Exception:
            return out

        for blip in blips:
            try:
                rid = blip.get("{%s}embed" % ns["r"])
                if not rid:
                    continue
                part = raw_doc.part.related_parts.get(rid)
                if part is None:
                    continue
                data = part.blob
                content_type = getattr(part, "content_type", "") or ""
                ext_hint = "." + content_type.split("/")[-1] if "/" in content_type else ""

                alt = None
                doc_pr = para._p.find(".//wp:docPr", ns)
                if doc_pr is not None:
                    alt = doc_pr.get("descr") or doc_pr.get("title")

                # process_image decodes+validates the bytes itself rather than
                # trusting content_type -- Word sometimes embeds vector
                # metafiles (WMF/EMF) that no browser can display, and a
                # content-type-only ext guess used to mislabel those as PNG,
                # producing a broken <img> and a useless payload sent for
                # captioning. Images that fail real decoding are skipped
                # rather than embedded broken.
                processed = process_image(data, ext_hint, existing_alt=alt)
                if processed is None:
                    continue
                image_counter[0] += 1
                caption = processed["caption"] or fallback_caption(image_counter[0], alt)

                out.append(Block(
                    id=make_id(), type=BlockType.IMAGE,
                    content=caption,
                    metadata={"src": processed["uri"], "ocr_text": processed["ocr_text"], "alt": alt or caption},
                ))
            except Exception:
                continue
        return out

    def _extract_runs(self, para) -> str:
        """
        Convert paragraph runs to inline markdown-ish text.
        Bold -> **text**, Italic -> *text*, both -> ***text***.
        Consecutive runs with identical formatting are merged into one
        wrapped span first: Word very commonly splits what visually reads
        as a single bold phrase into several runs (a manual line break
        inside the bold text, spell-check, copy/paste), and wrapping each
        one separately produced "****" between adjacent bold runs, which
        the markdown renderer downstream couldn't cleanly pair back up --
        leaving stray literal asterisks in the rendered text.

        A bold run that's also explicitly, deliberately colored (not just
        near-black default text) additionally gets wrapped in a
        {{c:HEXCOLOR}}...{{/c}} marker, preserved through to every export
        format. This is what makes documents like a script or transcript
        with color-coded speaker labels ("Simra:" in blue, "Aanya:" in
        maroon, each label its own short bold run followed by plain-colored
        body text) come through import looking like the original instead
        of collapsing to a single generic bold black label.
        """
        groups = []  # list of [bold, italic, text, color_or_None]
        for run in para.runs:
            text = run.text
            if not text:
                continue
            fmt = (bool(run.bold), bool(run.italic))
            color = None
            if run.bold and run.font.color and run.font.color.rgb:
                hexcolor = str(run.font.color.rgb)
                if not _is_grayscale(hexcolor):
                    color = hexcolor
            if groups and (groups[-1][0], groups[-1][1]) == fmt:
                groups[-1][2] += text
                # Only keep carrying a color forward if every run merged
                # into this group agrees on it -- a bold group that starts
                # colored and then continues in plain black shouldn't have
                # the whole merged span tagged with that first run's color.
                if groups[-1][3] != color:
                    groups[-1][3] = None
            else:
                groups.append([fmt[0], fmt[1], text, color])

        parts = []
        for bold, italic, text, color in groups:
            inner = text
            if color:
                inner = f"{{{{c:{color}}}}}{inner}{{{{/c}}}}"
            if bold and italic:
                parts.append(f"***{inner}***")
            elif bold:
                parts.append(f"**{inner}**")
            elif italic:
                parts.append(f"*{inner}*")
            else:
                parts.append(inner)
        return "".join(parts)

    def _strip_inline_md(self, text: str) -> str:
        """Remove inline markdown markers for use in headings."""
        text = re.sub(r"\{\{c:[0-9A-Fa-f]{6}\}\}|\{\{/c\}\}", "", text)
        return re.sub(r"\*+", "", text).strip()

    def _maybe_entity(self, text: str, level: int) -> BlockType:
        """Decide if a heading should be promoted to an entity block. Only
        promotes headings that actually look like a proper name (a
        worldbuilding character, place, faction...), not just any short
        heading that starts with a capital letter -- document scaffolding
        like "SUMMER ENGAGEMENT WORK", "SESSION 2026-27", or a field label
        ending in ":" all satisfy that looser test just as easily as a real
        name does."""
        text = text.strip()
        if level < 3 or not text or text.endswith(":"):
            return BlockType.HEADING
        words = text.split()
        if not words or len(words) > 5:
            return BlockType.HEADING
        if any(any(ch.isdigit() for ch in w) for w in words):
            return BlockType.HEADING
        # Real names aren't typeset in full block capitals in running
        # text; that's a much stronger signal of a section label.
        if any(w.isupper() and len(w) > 1 for w in words):
            return BlockType.HEADING
        if not all(w[0].isupper() for w in words if w[:1].isalpha()):
            return BlockType.HEADING
        generic = {
            "overview", "introduction", "summary", "background", "notes",
            "rules", "settings", "world", "timeline", "characters", "locations",
            "appendix", "references", "conclusion", "chapter", "section",
            "part", "contents", "index", "preface", "foreword",
            "submitted", "by", "session", "class", "roll", "aim", "objective",
            "objectives", "program", "output", "input", "declaration",
            "certificate", "acknowledgement", "acknowledgment", "engagement",
            "school", "college", "university", "department", "subject",
            "assignment", "project", "report", "practical", "viva",
            "examination", "number", "date", "name",
        }
        if any(w.lower().strip(":,&") in generic for w in words):
            return BlockType.HEADING
        return BlockType.ENTITY

    def _classify_paragraph(self, text: str, make_id) -> Optional[Block]:
        """Apply semantic pattern matching to plain paragraph text.

        Matches against the color-marker-stripped text -- a paragraph that
        starts with a colored+bold "Warning:"/"Note:"/etc. label (a very
        common real-world pattern) would otherwise fail to match any of
        these patterns, since the {{c:HEXCODE}} marker sits between the
        ** and the actual word. The returned Block still carries the
        original marker-preserving `text`, so the label keeps its color
        when rendered -- only the classification decision uses the
        stripped version.
        """
        plain = _strip_color_marker(text)
        m = WARNING_RE.match(plain)
        if m:
            return Block(id=make_id(), type=BlockType.WARNING, content=text)

        m = CALLOUT_RE.match(plain)
        if m:
            return Block(id=make_id(), type=BlockType.CALLOUT, content=text)

        m = RELATIONSHIP_RE.match(plain)
        if m:
            return Block(id=make_id(), type=BlockType.RELATIONSHIP, content=text)

        m = TIMELINE_RE.match(plain)
        if m:
            return Block(id=make_id(), type=BlockType.TIMELINE_EVENT, content=text)

        m = DEFINITION_RE.match(plain)
        if m:
            return Block(id=make_id(), type=BlockType.DEFINITION, content=text)

        # Blockquote: indented text (Word "Quote" / "Intense Quote" styles handled separately)
        return None

    def _get_outline_level(self, para) -> Optional[int]:
        """Read pPr/outlineLvl from paragraph XML."""
        try:
            from docx.oxml.ns import qn
            pPr = para._p.pPr
            if pPr is None:
                return None
            outlineLvl = pPr.find(qn("w:outlineLvl"))
            if outlineLvl is not None:
                val = outlineLvl.get(qn("w:val"))
                if val is not None:
                    return int(val)
        except Exception:
            pass
        return None

    def _is_list_paragraph(self, para) -> bool:
        """Detect list paragraphs by numPr XML element."""
        try:
            from docx.oxml.ns import qn
            pPr = para._p.pPr
            if pPr is None:
                return False
            return pPr.find(qn("w:numPr")) is not None
        except Exception:
            return False

    def _has_highlight(self, para) -> bool:
        """Check if any run has a highlight color (yellow = important)."""
        try:
            from docx.oxml.ns import qn
            for run in para.runs:
                rPr = run._r.rPr
                if rPr is not None:
                    highlight = rPr.find(qn("w:highlight"))
                    if highlight is not None:
                        return True
        except Exception:
            pass
        return False

    def _parse_table(self, tbl) -> str:
        """Convert a docx Table into pipe-delimited markdown table text."""
        rows = []
        try:
            for i, row in enumerate(tbl.rows):
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        except Exception:
            pass
        return "\n".join(rows)

    def _extract_title(self, blocks: List[Block], filepath: str) -> str:
        for b in blocks:
            if b.type == BlockType.HEADING and b.level == 1:
                return b.content
            if b.type == BlockType.ENTITY and b.level == 1:
                return b.content
        # Fallback to filename
        import os
        return os.path.splitext(os.path.basename(filepath))[0].replace("_", " ").title()

    def _build_sections(self, blocks: List[Block]) -> List[Section]:
        sections = []
        current = None
        counter = 0

        for block in blocks:
            is_h1 = block.type == BlockType.HEADING and block.level == 1
            is_h2 = block.type == BlockType.HEADING and block.level == 2

            if is_h1:
                if current:
                    sections.append(current)
                counter += 1
                current = Section(id=f"s{counter:03d}", title=block.content,
                                  level=1, blocks=[block])
            elif is_h2:
                if current is None:
                    counter += 1
                    current = Section(id=f"s{counter:03d}", title="Document", level=1)
                counter += 1
                sub = Section(id=f"s{counter:03d}", title=block.content,
                              level=2, blocks=[block])
                current.subsections.append(sub)
            else:
                if current is None:
                    counter += 1
                    current = Section(id=f"s{counter:03d}", title="Document", level=1)
                if current.subsections:
                    current.subsections[-1].blocks.append(block)
                else:
                    current.blocks.append(block)

        if current:
            sections.append(current)
        return sections
