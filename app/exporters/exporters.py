from app.models.document import Document
from app.renderer import decorations


def _parse_table_rows(raw: str) -> list:
    """Parse pipe-delimited markdown table text into a list of row cell-lists,
    skipping the '| --- | --- |' separator row."""
    import re
    rows = []
    for line in raw.strip().splitlines():
        line = line.strip()
        if not line or re.match(r"^\|[-\s|:]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if cells:
            rows.append(cells)
    return rows


class PDFExporter:
    """Exports rendered HTML to PDF using WeasyPrint."""

    def export(self, html: str, output_path: str) -> str:
        try:
            from weasyprint import HTML
            HTML(string=html).write_pdf(output_path)
            return output_path
        except ImportError:
            raise RuntimeError("WeasyPrint not installed. Run: pip install weasyprint")
        except Exception as e:
            raise RuntimeError(f"PDF export failed: {e}")


def _is_manuscript(doc: Document) -> bool:
    """Prose-heavy documents (novels, narrative nonfiction) read far better
    with book typesetting than the dense note-taking layout built for
    structured docs with lots of entities/callouts/tables."""
    from app.models.document import BlockType
    blocks = doc.all_blocks
    if len(blocks) < 20:
        return False
    prose_types = {BlockType.PARAGRAPH, BlockType.HEADING, BlockType.QUOTE, BlockType.DIVIDER}
    prose_count = sum(1 for b in blocks if b.type in prose_types)
    return prose_count / len(blocks) > 0.9


def _add_page_border(section, color_hex: str):
    """Word has a real page-border feature (Design > Page Borders) but
    python-docx has no high-level API for it -- inject the underlying
    <w:pgBorders> OOXML directly."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    sectPr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "18")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), color_hex.lstrip("#").upper())
        pg_borders.append(el)
    sectPr.append(pg_borders)


def _add_margin_ornaments(section, top_text: str, bottom_text: str, color_hex: str, rgb_fn):
    """Place a centered ornamental rule in the header (top margin) and a
    matching one in the footer (bottom margin) -- headers/footers repeat
    on every page automatically in Word, giving real per-page margin
    decoration rather than a one-time decoration."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for container, text in ((section.header, top_text), (section.footer, bottom_text)):
        para = container.paragraphs[0]
        para.text = ""
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(13)
        run.font.color.rgb = rgb_fn(color_hex)


class DOCXExporter:
    """Exports a Document model to a DOCX file.

    If `palette` is provided (a dict from auto_designer.to_docx_palette),
    colors and fonts are pulled from it instead of the hardcoded defaults
    below, so DOCX export can reflect the auto-generated per-document theme
    the same way PPTX export already does.

    Prose-heavy documents (novels) automatically get book-style typesetting:
    a fresh page per chapter and indented paragraphs instead of the default
    note-taking layout.
    """

    # Hardcoded fallback palette, used when no auto-design palette is passed
    # in (fixed presets, or theme couldn't be resolved).
    _DEFAULT_PALETTE = {
        "accent": "#2B4C9B",
        "text": "#000000",
        "entity": "#5D3A8E",
        "callout": "#1A6E4A",
        "warning": "#C0392B",
        "border": "#888888",
        "heading_font": None,
        "body_font": None,
    }

    def export(self, doc: Document, output_path: str, palette: dict = None,
               decorate: bool = False, border: bool = False,
               decoration_style: str = "auto", theme: str = None) -> str:
        try:
            from docx import Document as DocxDocument
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("python-docx not installed. Run: pip install python-docx")

        from app.models.document import BlockType

        manuscript = _is_manuscript(doc)
        pal = {**self._DEFAULT_PALETTE, **(palette or {})}
        if manuscript and not pal["body_font"]:
            pal = {**pal, "heading_font": pal["heading_font"] or "Times New Roman",
                   "body_font": "Times New Roman"}

        def rgb(hex_color: str) -> "RGBColor":
            h = hex_color.lstrip("#")
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        d = DocxDocument()

        # Base document font, if the auto-design specified one
        if pal["body_font"]:
            normal = d.styles["Normal"]
            normal.font.name = pal["body_font"]
        if manuscript:
            d.styles["Normal"].font.size = Pt(11.5)

        # Title
        title_para = d.add_heading(doc.title, 0)
        title_para.runs[0].font.color.rgb = rgb(pal["accent"])
        if pal["heading_font"]:
            title_para.runs[0].font.name = pal["heading_font"]
        if manuscript:
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if doc.project_type and doc.project_type.value != "Document" and not manuscript:
            meta = d.add_paragraph(f"{doc.project_type.value} — {len(doc.all_blocks)} blocks")
            meta.runs[0].font.size = Pt(9)
            meta.runs[0].font.color.rgb = rgb(pal["border"])

        if manuscript:
            d.add_page_break()

        first_h1_seen = False
        for block in doc.all_blocks:
            if manuscript and block.type == BlockType.HEADING and block.level == 1:
                if first_h1_seen:
                    d.add_page_break()
                first_h1_seen = True

            if block.type == BlockType.HEADING:
                level = min(block.level, 9)
                h = d.add_heading(block.content, level)
                if h.runs:
                    h.runs[0].font.color.rgb = rgb(pal["accent"])
                    if pal["heading_font"]:
                        h.runs[0].font.name = pal["heading_font"]
                if manuscript:
                    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif block.type == BlockType.PARAGRAPH:
                p = d.add_paragraph(block.content)
                if pal["body_font"] and p.runs:
                    p.runs[0].font.name = pal["body_font"]
                if manuscript:
                    p.paragraph_format.first_line_indent = Pt(18)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.3

            elif block.type == BlockType.IMAGE:
                self._add_image(d, block, pal)

            elif block.type == BlockType.ENTITY:
                p = d.add_paragraph()
                run = p.add_run(f"◆ {block.content}")
                run.bold = True
                run.font.color.rgb = rgb(pal["entity"])

            elif block.type in (BlockType.CALLOUT, BlockType.WARNING):
                prefix = "! WARNING: " if block.type == BlockType.WARNING else "NOTE: "
                p = d.add_paragraph(prefix + block.content)
                p.runs[0].bold = True
                if block.type == BlockType.WARNING:
                    p.runs[0].font.color.rgb = rgb(pal["warning"])
                else:
                    p.runs[0].font.color.rgb = rgb(pal["callout"])

            elif block.type == BlockType.TIMELINE_EVENT:
                p = d.add_paragraph(f"→ {block.content}", style="List Bullet")

            elif block.type == BlockType.RELATIONSHIP:
                p = d.add_paragraph(block.content)
                p.runs[0].italic = True

            elif block.type == BlockType.QUOTE:
                p = d.add_paragraph(f'"{block.content}"')
                p.runs[0].italic = True
                p.runs[0].font.color.rgb = rgb(pal["entity"])

            elif block.type == BlockType.LIST:
                for item in block.content.split("\n"):
                    if item.strip():
                        d.add_paragraph(item.strip(), style="List Bullet")

            elif block.type == BlockType.TASK:
                self._add_task(d, block.content, bool(block.metadata.get("checked")),
                                block.metadata.get("subtasks") or [], depth=0)

            elif block.type == BlockType.CODE:
                p = d.add_paragraph(block.content)
                p.runs[0].font.name = "Courier New"
                p.runs[0].font.size = Pt(9)

            elif block.type == BlockType.TABLE:
                rows = _parse_table_rows(block.content)
                if rows:
                    n_cols = max(len(r) for r in rows)
                    table = d.add_table(rows=len(rows), cols=n_cols)
                    table.style = "Light Grid Accent 1"
                    for ri, row in enumerate(rows):
                        for ci in range(n_cols):
                            cell_text = row[ci] if ci < len(row) else ""
                            cell = table.cell(ri, ci)
                            cell.text = cell_text
                            if ri == 0:
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.bold = True
                    d.add_paragraph("")

            elif block.type == BlockType.DIVIDER:
                d.add_paragraph("─" * 40)

            else:
                if block.content.strip():
                    d.add_paragraph(block.content)

        if border or decorate:
            section = d.sections[0]
            if border:
                _add_page_border(section, pal["accent"])
            if decorate:
                resolved_style = decorations.resolve_style(decoration_style, doc, theme)
                primary, secondary = decorations.DECORATION_SETS.get(
                    resolved_style, decorations.DECORATION_SETS[decorations.DEFAULT_SET])
                top_text = decorations.margin_text(primary)
                bottom_text = decorations.margin_text(secondary)
                _add_margin_ornaments(section, top_text, bottom_text, pal["accent"], rgb)

        d.save(output_path)
        return output_path

    def _add_image(self, d, block, pal):
        """Embed an IMAGE block's data URI as an actual picture, sized to
        fit within the page margins, with an italic caption paragraph if
        one is available. Never raises -- a missing/undecodable image is
        skipped rather than failing the whole export."""
        import io
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from app.parser.image_pipeline import from_data_uri, is_data_uri

        meta = block.metadata or {}
        src = meta.get("src") or block.content or ""
        caption = meta.get("caption") or ""
        if block.content and block.content != src and not caption:
            caption = block.content

        if not is_data_uri(src):
            # External URL or unresolvable reference -- nothing we can embed
            # directly into the docx; skip rather than break the export.
            return
        data, _ext = from_data_uri(src)
        if not data:
            return

        try:
            max_width = Inches(6.0)
            width_arg = max_width
            try:
                from PIL import Image
                with Image.open(io.BytesIO(data)) as im:
                    px_w, dpi_w = im.width, (im.info.get("dpi", (96, 96))[0] or 96)
                    native_width = Inches(px_w / dpi_w)
                    if native_width < max_width:
                        width_arg = native_width
            except Exception:
                pass  # fall back to the fixed max width
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(io.BytesIO(data), width=width_arg)
        except Exception:
            return

        if caption:
            cap = d.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap.runs:
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9.5)
                cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def _add_task(self, d, text, checked, subtasks, depth):
        """Recursively add a task and its subtasks as indented checkbox lines."""
        from docx.shared import Pt, RGBColor
        box = "[x]" if checked else "[ ]"
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Pt(18 * depth)
        run = p.add_run(f"{box} {text}")
        if checked:
            run.font.strike = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for sub in subtasks:
            self._add_task(d, sub.get("text", ""), bool(sub.get("checked")),
                            sub.get("subtasks") or [], depth + 1)
